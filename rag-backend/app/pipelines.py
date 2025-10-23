# rag-backend/app/pipelines.py

from typing import List, Optional, Dict, Any, Iterable
import os
import io
import re
import tempfile
import pathlib

from haystack import Pipeline, Document
from haystack.components.preprocessors import DocumentCleaner, DocumentSplitter
from haystack.components.writers import DocumentWriter
from haystack.components.builders import PromptBuilder

# Built-in Haystack converters
from haystack.components.converters import (
    PyPDFToDocument,
    TextFileToDocument,
    MarkdownToDocument,
    HTMLToDocument,
)

# Abhängigkeiten für zusätzliche Formate
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from odf.opendocument import load as odf_load
from odf import text as odf_text
from openpyxl import load_workbook
import mailparser

from .deps import (
    get_document_store,
    get_doc_embedder,
    get_text_embedder,
    get_retriever,
    get_generator,
)


def int_env(k: str, d: int) -> int:
    try:
        return int(os.getenv(k, d))
    except Exception:
        return d


# -------------------------------
# Konvertierung: Bytes -> Documents
# -------------------------------

_PDF = PyPDFToDocument()
_TXT = TextFileToDocument()
_MD = MarkdownToDocument()
_HTML = HTMLToDocument()


def _doc_from_text(text: str, meta: Optional[Dict[str, Any]] = None) -> Document:
    cleaned = re.sub(r"\r\n?", "\n", text or "").strip()
    return Document(content=cleaned, meta=meta or {})


def _convert_docx(data: bytes, meta: Dict[str, Any]) -> List[Document]:
    f = io.BytesIO(data)
    doc = DocxDocument(f)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            cells = [c.text for c in row.cells]
            parts.append("\t".join(cells))
    text = "\n".join([p for p in parts if p])
    return [_doc_from_text(text, meta)] if text.strip() else []


def _extract_odf_text(elem) -> Iterable[str]:
    for node in elem.childNodes:
        if isinstance(node, odf_text.P) or isinstance(node, odf_text.H):
            yield "".join(t.data for t in node.childNodes if hasattr(t, "data"))
        elif isinstance(node, odf_text.List) or isinstance(node, odf_text.ListItem):
            yield from _extract_odf_text(node)
        else:
            if hasattr(node, "childNodes"):
                yield from _extract_odf_text(node)


def _convert_odt(data: bytes, meta: Dict[str, Any]) -> List[Document]:
    f = io.BytesIO(data)
    odoc = odf_load(f)
    texts = []
    for body in odoc.getElementsByType(odf_text.P):
        texts.append("".join(t.data for t in body.childNodes if hasattr(t, "data")))
    if not texts:
        try:
            body = odoc.text
            texts = list(_extract_odf_text(body))
        except Exception:
            texts = []
    text = "\n".join([t for t in texts if t])
    return [_doc_from_text(text, meta)] if text.strip() else []


def _convert_xlsx(data: bytes, meta: Dict[str, Any]) -> List[Document]:
    f = io.BytesIO(data)
    wb = load_workbook(f, read_only=True, data_only=True)
    docs: List[Document] = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else str(v) for v in row]
            rows.append("\t".join(vals))
        text = "\n".join(rows).strip()
        if text:
            m = dict(meta)
            m["sheet_name"] = ws.title
            docs.append(_doc_from_text(text, m))
    return docs


def _convert_eml(data: bytes, meta: Dict[str, Any]) -> List[Document]:
    mail = mailparser.parse_from_bytes(data)
    headers = {
        "from": (mail.from_[0][1] if mail.from_ else None),
        "to": [r[1] for r in (mail.to or [])],
        "subject": mail.subject,
        "date": str(mail.date) if mail.date else None,
        "message_id": mail.message_id,
    }
    body_text = ""
    if mail.text_plain:
        body_text = "\n\n".join(mail.text_plain)
    elif mail.text_html:
        texts = []
        for html in mail.text_html:
            soup = BeautifulSoup(html, "lxml")
            texts.append(soup.get_text(separator="\n"))
        body_text = "\n\n".join(texts)
    content = []
    content.append(f"Subject: {headers.get('subject') or ''}")
    if headers.get("from"):
        content.append(f"From: {headers['from']}")
    if headers.get("to"):
        content.append(f"To: {', '.join(headers['to'])}")
    if headers.get("date"):
        content.append(f"Date: {headers['date']}")
    content.append("")
    content.append(body_text or "")
    meta_all = dict(meta)
    meta_all.update({"email_headers": headers})
    return [_doc_from_text("\n".join(content), meta_all)]


def _run_converter_with_tempfile(converter, suffix: str, data: bytes) -> List[Document]:
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        out = converter.run(sources=[tmp_path])["documents"]
        return out
    finally:
        try:
            pathlib.Path(tmp_path).unlink(missing_ok=True)
        except Exception:
            pass


def convert_bytes_to_documents(
    filename: str,
    mime: str,
    data: bytes,
    default_meta: Optional[Dict[str, Any]] = None
) -> List[Document]:
    """
    Universeller Konverter: Datei-Bytes -> Haystack Documents
    Unterstützte Typen: txt, md, pdf, html, docx, odt, xlsx, eml
    """
    meta = dict(default_meta or {})
    meta.update({"filename": filename, "mime": mime})

    mime = (mime or "").lower()

    try:
        if mime == "application/pdf" or filename.lower().endswith(".pdf"):
            docs = _run_converter_with_tempfile(_PDF, ".pdf", data)

        elif mime in ("text/markdown",) or filename.lower().endswith(".md"):
            docs = _run_converter_with_tempfile(_MD, ".md", data)

        elif mime in ("text/html", "application/xhtml+xml") or filename.lower().endswith((".htm", ".html")):
            docs = _run_converter_with_tempfile(_HTML, ".html", data)

        elif mime in ("text/plain",) or filename.lower().endswith(".txt"):
            docs = _run_converter_with_tempfile(_TXT, ".txt", data)

        elif mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",) or filename.lower().endswith(".docx"):
            docs = _convert_docx(data, meta)

        elif mime in ("application/vnd.oasis.opendocument.text", "application/odt") or filename.lower().endswith(".odt"):
            docs = _convert_odt(data, meta)

        elif mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",) or filename.lower().endswith(".xlsx"):
            docs = _convert_xlsx(data, meta)

        elif mime in ("message/rfc822", "application/eml") or filename.lower().endswith(".eml"):
            docs = _convert_eml(data, meta)

        else:
            try:
                text = data.decode("utf-8", errors="ignore")
            except Exception:
                text = ""
            docs = [_doc_from_text(text, meta)] if text.strip() else []

        for d in docs:
            m = dict(d.meta or {})
            m.setdefault("filename", filename)
            m.setdefault("mime", mime)
            d.meta = m

        return docs

    except Exception as e:
        fallback = f"[Konvertierung fehlgeschlagen: {type(e).__name__}: {e}]"
        return [_doc_from_text(fallback, meta)]


# -------------------------------
# Pipelines
# -------------------------------

def build_index_pipeline():
    """Indexierungs-Pipeline: Dokumente -> Embeddings -> Qdrant"""
    store = get_document_store()
    embedder = get_doc_embedder()
    writer = DocumentWriter(document_store=store)

    cleaner = DocumentCleaner()
    splitter = DocumentSplitter(
        split_by="word",
        split_length=int_env("CHUNK_SIZE", 200),
        split_overlap=int_env("CHUNK_OVERLAP", 20),
    )

    pipe = Pipeline()
    pipe.add_component("clean", cleaner)
    pipe.add_component("split", splitter)
    pipe.add_component("embed", embedder)
    pipe.add_component("write", writer)

    pipe.connect("clean.documents", "split.documents")
    pipe.connect("split.documents", "embed.documents")
    pipe.connect("embed.documents", "write.documents")
    
    return pipe, store


def build_query_pipeline(store=None):
    """Query-Pipeline: Frage -> Retrieval -> Generation mit vLLM"""
    store = store or get_document_store()
    retriever = get_retriever(store)
    gen = get_generator()
    qembed = get_text_embedder()

    template = """Beantworte prägnant und korrekt anhand der folgenden Dokumente.
Gib keine Inhalte wieder, die nicht im Kontext stehen.

Kontext:
{% for d in documents %}
- {{ d.content | truncate(600) }}
{% endfor %}

Frage: {{ query }}
Antwort:"""

    pipe = Pipeline()
    pipe.add_component("embed_query", qembed)
    pipe.add_component("retrieve", retriever)
    pipe.add_component("prompt_builder", PromptBuilder(template=template))
    pipe.add_component("generate", gen)

    # Query-Text -> Embedder -> Retriever
    pipe.connect("embed_query.embedding", "retrieve.query_embedding")
    # Retriever-Dokumente -> PromptBuilder -> Generator
    pipe.connect("retrieve.documents", "prompt_builder.documents")
    pipe.connect("prompt_builder.prompt", "generate.prompt")
    
    return pipe